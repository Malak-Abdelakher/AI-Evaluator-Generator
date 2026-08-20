from pathlib import Path

import pymupdf
import pytest
from docx import Document as DocxDocument

from ingestion.base.loader_registry import get_file_loader
from ingestion.loaders.code_loader import CodeKnowledgeLoader
from ingestion.loaders.docx_loader import DOCXKnowledgeLoader
from ingestion.loaders.pdf_loader import PDFKnowledgeLoader
from ingestion.loaders.txt_loader import TXTKnowledgeLoader
from schemas.document import validate_document_metadata


def test_code_loader(tmp_path):
    file_path = tmp_path / "example.py"

    file_path.write_text(
        "def hello():\n    return 'Hello AI'",
        encoding="utf-8",
    )

    loader = CodeKnowledgeLoader()
    documents = loader.load(file_path)

    assert len(documents) == 1
    assert "def hello" in documents[0].page_content
    assert documents[0].metadata["source_type"] == "code"
    assert validate_document_metadata(documents[0].metadata)


def test_docx_loader(tmp_path):
    file_path = tmp_path / "example.docx"

    docx = DocxDocument()
    docx.add_heading("Artificial Intelligence", level=1)
    docx.add_paragraph(
        "LangChain can be used to build LLM applications."
    )
    docx.save(file_path)

    loader = DOCXKnowledgeLoader()
    documents = loader.load(file_path)

    assert len(documents) == 1
    assert "Artificial Intelligence" in documents[0].page_content
    assert "LangChain" in documents[0].page_content
    assert documents[0].metadata["source_type"] == "docx"
    assert validate_document_metadata(documents[0].metadata)


def test_pdf_loader(tmp_path):
    file_path = tmp_path / "example.pdf"

    pdf = pymupdf.open()
    page = pdf.new_page()

    page.insert_text(
        (72, 72),
        "LCEL connects runnable components together.",
    )

    pdf.save(file_path)
    pdf.close()

    loader = PDFKnowledgeLoader()
    documents = loader.load(file_path)

    assert len(documents) == 1
    assert "LCEL" in documents[0].page_content
    assert documents[0].metadata["source_type"] == "pdf"
    assert documents[0].metadata["page"] == 1
    assert validate_document_metadata(documents[0].metadata)


def test_registry_resolves_txt():
    loader = get_file_loader("example.txt")

    assert isinstance(loader, TXTKnowledgeLoader)


def test_registry_resolves_pdf():
    loader = get_file_loader("example.pdf")

    assert isinstance(loader, PDFKnowledgeLoader)


def test_registry_resolves_docx():
    loader = get_file_loader("example.docx")

    assert isinstance(loader, DOCXKnowledgeLoader)


def test_registry_resolves_code():
    loader = get_file_loader("example.py")

    assert isinstance(loader, CodeKnowledgeLoader)


def test_registry_rejects_unknown_file_type():
    with pytest.raises(ValueError):
        get_file_loader("example.xyz")