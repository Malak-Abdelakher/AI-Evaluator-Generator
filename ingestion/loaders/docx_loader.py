from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document

from ingestion.base.base_loader import BaseKnowledgeLoader


class DOCXKnowledgeLoader(BaseKnowledgeLoader):
    """Loader for DOCX knowledge sources."""

    def load(self, source: str | Path) -> list[Document]:
        path = Path(source)

        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")

        if not path.is_file():
            raise ValueError(f"DOCX source is not a file: {path}")

        if path.suffix.lower() != ".docx":
            raise ValueError(
                f"DOCXKnowledgeLoader only accepts .docx files: {path}"
            )

        docx = DocxDocument(path)

        content_parts = []

        # Extract normal paragraphs.
        for paragraph in docx.paragraphs:
            text = paragraph.text.strip()

            if text:
                content_parts.append(text)

        # Extract tables as textual rows.
        for table in docx.tables:
            for row in table.rows:
                cells = [
                    cell.text.strip()
                    for cell in row.cells
                ]

                row_text = " | ".join(cells)

                if row_text.strip():
                    content_parts.append(row_text)

        content = "\n".join(content_parts)

        return [
            Document(
                page_content=content,
                metadata={
                    "source": str(path),
                    "source_type": "docx",
                    "filename": path.name,
                },
            )
        ]